# -*- coding: utf-8 -*-
"""
Created on Tue Oct  9 09:33:50 2018

@author: xmlon
"""
from docx import Document
import os, requests
from translate import translate
from time import sleep
from random import randint

class Word2word():
#    word -> translate -> word -> pdf
#    initialize the object with a word docx file and its name
    def __init__(self, source_word):
        cwd = os.getcwd()
        word_path = cwd + '\\' + source_word
        self.word_document = Document(word_path)
        self.word_document_zh = Document(word_path)
        self.word_name = source_word
        self.state_paragraphs = "hasn't started"
        self.state_tables = "hasn't started"
    # translate the word document
    # translate paragraphs first
    def list_text(self):
        len_doc = len(self.word_document.paragraphs)
        text_list = []
        for i in range(len_doc):
            text = self.word_document.paragraphs[i].text
            # list those having texts, exclude those having no texts
            if text.strip() != '':
                text_list.append([i,text+'\n'])
                # '\n' for correct words counting
        self.text_list = text_list
        self.zh_list = []

    # split text_list into different parts each of which has no more
    # words than 4800 in total
    def get_bounds_text_list(self):
        bounds = []
        len_list = len(self.text_list)
        if len_list == 0:
            print('There is no paragraphs containing text')
            return()
        # self.text_list[i] = [index_of_paragraph, its_text]
        lengths = [len(l[1]) for l in self.text_list]
        n = 1
        s = 0 # l for the starting index
        GTL = 4800 # GTL: Google Translate Length
        while n <= len_list:
            # I have ignored the possibility that the length of a
            # single paragraph may exceed GTL(4800) words
            if sum(lengths[s:n])<=GTL and sum(lengths[s:(n+1)])>GTL:
                s = n
                bounds.append(n)
            n = n + 1
        # checking whether the index of the last paragraph is
        # included or not
        if len(bounds) >= 1:
            if bounds[-1] != len_list:
                bounds.append(len_list)
        else:
            bounds.append(len_list)

        self.text_list_bounds = bounds

    # combine each part into a whole string,
    # translate the string,
    # split the translated string into a list,
    # and populate self.zh_list
    def combine_translate_split(self,s,e):
        # combine the texts into a text string for translation
        text_string = ''
        for i in range(s,e):
            text_string = text_string + self.text_list[i][1]
            # self.text_list[i] = [index_of_paragraph, its_text]
        # tranlsate the string
        Failed = True
        try:
            translated_string = translate(text_string)
            print('successfully translated: '+translated_string)
            Failed = False
        except:
            print('Something wrong with result = translate(text)',
                  ' and the text is: '+text_string)
        if Failed:
            return()
        # split the translated_string into a list
        # and populate it into self.zh_list
        translated_list = translated_string.split('\n')
        # 'I\nlike\n'.split('\n') = ['I', 'like', '']
        # But translate('I\nlike\n').split('\n') = ['我', '喜欢']
        if len(translated_list) != (e-s):
            print("The return translated_string doesn't have ",
                  "the same number of paragraphs as text_string!")
        for i in range(s,e):
            index = self.text_list[i][0]
            text = translated_list[i-s]
            self.zh_list.append([index,text])

    # translate all the paragraphs and populate self.zh_list
    def translate_all_paragraphs(self):
        s = 0
        bounds = self.text_list_bounds
        for i in range(len(bounds)):
            e = bounds[i]
            self.combine_translate_split(s,e)
            s = e
            sleep(randint(10,20))
            txt ="translated to %s th element of word2pdf.text_list"
            self.state_paragraphs = txt%str(e)
        self.translate_state = 'translated all paragraphs'

    # substitute all the translated texts back into their
    # corresponding paragraphs
    def substitute_paragraphs(self):
        # substitute self.word_document_zh.paragraphs
        for i in range(len(self.zh_list)):
            index,text = self.zh_list[i][:2]
            self.word_document_zh.paragraphs[index].text = text

#    translate tables in word_document
    def list_table_text(self):
        # combine the texts in cells into
        # a list [[n,[r,c,text],...],...]
        tables = self.word_document.tables
        len_tables = len(tables)
        table_text_list = []
        table_text_list_zh = []
        for n in range(len_tables):
            table = tables[n]
            table_text_list.append([n,[]])
            table_text_list_zh.append([n,[]])
            try:
                for i in range(len(table.rows)):
                    for j in range(len(table.rows[i].cells)):
                        cell = table.rows[i].cells[j]
                        text = cell.text
                        if text.strip() != '':
                            table_text_list[n][1].append([i,j,
                                           text + '\n'])
                            table_text_list_zh[n][1].append([i,j,''])
            except:
                print('of table %s'%n)

        self.table_text_list = table_text_list
        self.table_text_list_zh = table_text_list_zh

    # I ignored the possibility that the words length of
    # total texts in all cells of each table may exceed
    # Google translate length limits 4800
    def combine_translate_split_table(self,i_table):
        # combine words together
        text_string = ''
        table_text_list = self.table_text_list[i_table]
        cells_list = table_text_list[1]
        # table_text_list = [n,cells_list]
        if len(cells_list) == 0:
            return()
        for i in range(len(cells_list)):
            cell = cells_list[i] # cells_list = [cell, cell, ...]
            text_string = text_string + cell[2] # cell =[r,c,text]
        # translate
        Failed = True
        try:
            translated_string = translate(text_string)
            print('successfully translated: %s'%text_string)
            Failed = False
        except:
            print('Something wrong with result = translate(text)',
                  ' and the text is: %s'%text_string)
        if Failed:
            return()
        # split the translated texts into zh (translated) list
        table_text_list_zh = self.table_text_list_zh[i_table]
        cells_list_zh = table_text_list_zh[1]
        translated_list = translated_string.split('\n')
        for i in range(len(cells_list_zh)):
            cells_list_zh[i][2] = translated_list[i]

    def translate_all_tables(self):
        len_tables = len(self.word_document.tables)
        for i in range(len_tables):
            len_of_text = len(self.table_text_list[i][1])
            # self.table_text_list[i] = [i,[...]]
            # if the table has no texts in its cells, ignore it
            self.state_tables = "translated to %s table"%str(i)
            if len_of_text != 0:
                self.combine_translate_split_table(i)
                sleep(randint(10,15))
        self.state_tables = "translated all tables"

    def substitute_tables(self):
        # substitute one table by one table
        len_tables = len(self.word_document.tables)
        for n in range(len_tables):
            table = self.word_document_zh.tables[n]
            cells = self.table_text_list_zh[n][1]
            len_of_cells = len(cells)
            if len_of_cells != 0:
                for i in range(len_of_cells):
                    r,c,text = cells[i]
                    table.rows[r].cells[c].text = text

    def translate_paragraphs_in_one_shot(self):
        self.list_text()
        self.get_bounds_text_list()
        self.translate_all_paragraphs()
        self.substitute_paragraphs()

    def translate_tables_in_one_shot(self):
        self.list_table_text()
        self.translate_all_tables()
        self.substitute_tables()

    def translate_all_in_one_shot(self):
        self.translate_paragraphs_in_one_shot()
        if len(self.word_document.tables) >= 1:
            self.translate_tables_in_one_shot()

    def save2docx(self):
        self.word_document_zh.save(self.word_name[:-5]+'_中文.docx')

#if __name__ == '__main__':
#    import sys
#    word = Word2word(sys.argv[1])
#    word.translate_all_in_one_shot()
#    word.save2docx()
